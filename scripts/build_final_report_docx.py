from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "final_report_grade9.md"
OUTPUT = ROOT / "docs" / "ProcureWise_Final_Report.docx"

TEXT = RGBColor(0, 0, 0)


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_separate, text, fld_char_end])


def apply_mla_run_style(paragraph) -> None:
    paragraph.paragraph_format.line_spacing = 2.0
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.color.rgb = TEXT


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 2.0

    for name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.font.color.rgb = TEXT
        style.font.bold = True
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 2.0

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 2.0


def add_header(document: Document) -> None:
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("Ferrufino ")
    add_field(paragraph, "PAGE")
    apply_mla_run_style(paragraph)


def add_mla_heading(document: Document) -> None:
    for text in [
        "Jesus Arteaga, Omar Ferrufino, and Marie Trejo Sandoval",
        "Professor Sam Gill",
        "ISYS 573-01",
        "10 May 2026",
    ]:
        paragraph = document.add_paragraph(text)
        apply_mla_run_style(paragraph)

    title = document.add_paragraph("ProcureWise AI Agent Final Report")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_mla_run_style(title)


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "")
    return text


def add_code_paragraph(document: Document, text: str) -> None:
    for code_line in text.splitlines() or [""]:
        paragraph = document.add_paragraph(code_line)
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.line_spacing = 2.0
        for run in paragraph.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(10.5)
            run.font.color.rgb = TEXT


def add_bold_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(clean_inline(text))
    run.bold = True
    apply_mla_run_style(paragraph)


def add_markdown_content(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    in_code = False
    code_lines: list[str] = []
    skip_until_exec = False

    for line in lines:
        stripped = line.strip()
        if stripped == "# ProcureWise AI Agent Final Report":
            skip_until_exec = True
            continue
        if skip_until_exec:
            if stripped == "## Executive Summary":
                skip_until_exec = False
            else:
                continue

        if stripped.startswith("```"):
            if in_code:
                add_code_paragraph(document, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not stripped:
            continue

        if stripped.startswith("## "):
            add_bold_heading(document, stripped[3:])
            continue
        if stripped.startswith("### "):
            add_bold_heading(document, stripped[4:])
            continue
        if stripped.startswith("#### "):
            add_bold_heading(document, stripped[5:])
            continue

        bullet = re.match(r"^- (.*)", stripped)
        if bullet:
            paragraph = document.add_paragraph(clean_inline(bullet.group(1)), style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            apply_mla_run_style(paragraph)
            continue

        numbered = re.match(r"^\d+\. (.*)", stripped)
        if numbered:
            paragraph = document.add_paragraph(clean_inline(numbered.group(1)), style="List Number")
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            apply_mla_run_style(paragraph)
            continue

        paragraph = document.add_paragraph(clean_inline(stripped))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        apply_mla_run_style(paragraph)


def main() -> None:
    document = Document()
    style_document(document)
    add_header(document)
    add_mla_heading(document)
    add_markdown_content(document, SOURCE.read_text(encoding="utf-8"))
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
